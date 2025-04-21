# filename: main.py

import asyncio
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from num2words import num2words
from pydantic import BaseModel, Field

# Pydantic AI imports
from pydantic_ai import Agent
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.providers.openai import OpenAIProvider

# Docx generation imports (will be used inside the MCP Python execution)
# Ensure these are importable in the environment where mcp-run-python runs
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    # These are not strictly needed in the main script's environment,
    # but helps with clarity / local testing if desired.
    print("INFO: 'python-docx' library not found in the main environment. "
          "Ensure it's available for the '@pydantic/mcp-run-python' server.")
    Document = None # type: ignore

# Load environment variables from .env file
load_dotenv(override=True)

# --- Configuration ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY not found in environment variables. Please set it in your .env file.")

# *** IMPORTANT: Update this path to the directory where your files are located ***
# This path will be used by the MCP filesystem server
BASE_WORKING_DIR = Path("/app/emanuel") # ADJUST THIS PATH

TEMPLATE_FILENAME = "Predložak Dodatka  ugovoru _ UVEĆANA UPLATA_DJELOMIČNA OTPLATA KREDITA _bez izmjene roka _ kraći rok.docx"
OUTPUT_FILENAME = "POPUNJEN_Dodatak_ugovoru_Barbara_Stazic.docx"

# Source documents (relative to BASE_WORKING_DIR) - We won't read them directly here,
# data is pre-extracted for MVP, but listing for completeness.
# SOURCE_PDF_STATUS = "DocumentServlet.pdf"
# SOURCE_PDF_CONTRACT = "FileDownloadServlet.pdf"
# SOURCE_PDF_REPAYMENT_PLAN = "Otplatni_plan.pdf"

# --- Pre-extracted Data for MVP ---
# (Based on analysis of provided PDFs: DocumentServlet.pdf, FileDownloadServlet.pdf, Otplatni_plan.pdf)
extracted_data = {
    "dodatak_datum_zakljucenja": "24.02.2025.",
    "dodatak_mjesto_zakljucenja": "ĐAKOVO",
    "dodatak_broj": "1",
    "korisnik_ime_prezime": "BARBARA STAŽIĆ",
    "korisnik_adresa": "JOSIPA JURJA STROSSMAYERA 129, 31000 OSIJEK",
    "korisnik_oib": "61897713961",
    "original_ugovor_naziv": "UGOVOR O NENAMJENSKOM GOTOVINSKOM KREDITU",
    "original_ugovor_broj_partije": "9919479387",
    "original_ugovor_datum_zakljucenja": "18.03.2024.",
    "original_ugovor_iznos": 11000.00,
    "original_ugovor_valuta": "EUR",
    "uplata_iznos": 3000.00, # Calculated: 9158.10 (from DocumentServlet) - 6158.10 (from Otplatni_plan)
    "nova_glavnica_iznos": 6158.10,
    "novi_anuitet_iznos": 185.26,
    "solidarni_duznik_ime_prezime": None, # Not found in docs
    "solidarni_duznik_adresa": None,
    "solidarni_duznik_oib": None,
    "solidarni_jamac_ime_prezime": None, # Not found in docs
    "solidarni_jamac_adresa": None,
    "solidarni_jamac_oib": None,
    # Add more fields as needed for other placeholders if missed
}

# --- Helper Function for Croatian Number Words ---
def num_to_hr_words(number, currency="EUR"):
    """Converts a number to Croatian words, handling euros and cents."""
    try:
        # Ensure num2words is available
        import num2words
    except ImportError:
        raise ImportError("The 'num2words' library is required. Please install it (`pip install num2words`).") from None

    if not isinstance(number, (int, float)):
        return str(number) # Return as string if not a number

    integer_part = int(number)
    fractional_part = round((number - integer_part) * 100)

    words = num2words.num2words(integer_part, lang='hr')

    if currency == "EUR":
        words += " eura"
        if fractional_part > 0:
            words += " i " + num2words.num2words(fractional_part, lang='hr') + " centi"
    elif currency == "HRK": # Example if needed later
         words += " kuna"
         if fractional_part > 0:
            words += " i " + num2words.num2words(fractional_part, lang='hr') + " lipa"
    else: # Handle non-currency numbers or other currencies
        if fractional_part > 0:
             # Generic decimal representation if not EUR/HRK
             words += " cijelih " + num2words.num2words(fractional_part, lang='hr') + " stotinki"


    return words

def date_to_hr_words(date_str):
    """ Converts a DD.MM.GGGG date string to Croatian words. """
    try:
        dt_obj = datetime.strptime(date_str, "%d.%m.%Y")
        day = dt_obj.day
        month = dt_obj.month
        year = dt_obj.year

        day_words = num2words.num2words(day, lang='hr', to='ordinal') # e.g., "prvi", "drugi" - check if genitive needed
        # Need month names in genitive case
        months_genitive = [
            "siječnja", "veljače", "ožujka", "travnja", "svibnja", "lipnja",
            "srpnja", "kolovoza", "rujna", "listopada", "studenog", "prosinca"
        ]
        month_words = months_genitive[month - 1]
        year_words = num2words.num2words(year, lang='hr') # e.g., "dvije tisuće dvadesetčetvrte"

        # Construct the phrase - requires careful grammar for ordinal day + month genitive
        # Simple construction for now, might need refinement for perfect grammar
        # Example target: "osamnaestog ožujka dvije tisuće dvadesetčetvrte"
        # Need mapping from ordinal word (e.g., "osamnaesti") to genitive ("osamnaestog")
        # This is complex, returning simple version for MVP
        day_ord_word = num2words.num2words(day, lang='hr', to='ordinal')
        # Basic adjustment for common cases (needs proper grammatical handling)
        if day_ord_word.endswith('i'):
            day_genitive_word = day_ord_word[:-1] + 'og' # prvi -> prvog, drugi -> drugog etc.
        else:
             day_genitive_word = day_ord_word # Approximation

        return f"{day_genitive_word} {month_words} {year_words}"

    except Exception as e:
        print(f"Error converting date {date_str} to words: {e}")
        return f"[Greška pri pretvaranju datuma {date_str}]"


# --- Pydantic Model for the Tool Input ---
class FillDocxArgs(BaseModel):
    template_file_path: str = Field(..., description="Relative path to the DOCX template file within the MCP filesystem.")
    output_file_path: str = Field(..., description="Relative path for the generated DOCX file within the MCP filesystem.")
    data: dict = Field(..., description="Dictionary containing data to fill into the template placeholders.")

# --- Python Code to run via MCP ---
# This code will be executed by the '@pydantic/mcp-run-python' server
# It needs access to 'python-docx' and 'num2words' libraries.

PYTHON_CODE_FOR_MCP = """
import sys
# Add parent directory to path if needed to find local helper functions/libs
# sys.path.append('..') # Uncomment if num_to_hr_words is in a separate file

from docx import Document
from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Uncomment if needed
import re
from datetime import datetime

try:
    # Attempt to import num2words; it must be installed in the MCP python env
    import num2words
except ImportError:
    print("ERROR: 'num2words' library not found in MCP Python environment. Install it.", file=sys.stderr)
    # Define a fallback function if num2words is missing
    def num_to_hr_words(number, currency="EUR"):
        return f"[num2words Greška: {number} {currency}]"
    def date_to_hr_words(date_str):
         return f"[num2words Greška: {date_str}]"
else:
    # Define the helper functions again within this execution context
    def num_to_hr_words(number, currency="EUR"):
        # Simplified version for brevity, copy the full one if needed
        if not isinstance(number, (int, float)): return str(number)
        integer_part = int(number)
        fractional_part = round((number - integer_part) * 100)
        words = num2words.num2words(integer_part, lang='hr')
        if currency == "EUR": words += " eura"
        else: words += f" {currency}" # Basic fallback
        if fractional_part > 0:
            words += " i " + num2words.num2words(fractional_part, lang='hr')
            if currency == "EUR": words += " centi"
            else: words += " stotinki" # Basic fallback
        return words

    def date_to_hr_words(date_str):
        # Simplified version for brevity, copy the full one if needed
        try:
            dt_obj = datetime.strptime(date_str, "%d.%m.%Y")
            day, month, year = dt_obj.day, dt_obj.month, dt_obj.year
            months_genitive = ["siječnja", "veljače", "ožujka", "travnja", "svibnja", "lipnja",
                               "srpnja", "kolovoza", "rujna", "listopada", "studenog", "prosinca"]
            day_ord = num2words.num2words(day, lang='hr', to='ordinal')
            # Very basic genitive adjustment
            day_gen = day_ord[:-1] + 'og' if day_ord.endswith('i') else day_ord
            return f"{day_gen} {months_genitive[month - 1]} {num2words.num2words(year, lang='hr')}"
        except Exception:
            return f"[Datum Greška: {date_str}]"


def fill_template(template_path, output_path, data):
    print(f"Attempting to fill template: {template_path}")
    print(f"Output path: {output_path}")
    print(f"Data received: {data}")

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"ERROR: Failed to open template file '{template_path}'. Make sure the path is correct and accessible by the filesystem server. Error: {e}", file=sys.stderr)
        return {"success": False, "error": f"Failed to open template: {e}"}

    # Prepare replacements, including calculated/formatted values
    replacements = {
        # Direct data
        '\\[DD\\.MM\\.GGGG\\.\\]': data.get('dodatak_datum_zakljucenja', '[DATUM_DODATKA]'), # Date for addendum signing
        '\\[upisati mjesto\\]': data.get('dodatak_mjesto_zakljucenja', '[MJESTO_DODATKA]'),
        'Dodatak br  \\. __': f"Dodatak br. {data.get('dodatak_broj', 'X')}", # Addendum number
        'Dodatak br\\.__': f"Dodatak br. {data.get('dodatak_broj', 'X')}", # Variations in placeholder
        'Dodatka br\\.__': f"Dodatka br. {data.get('dodatak_broj', 'X')}",
        'Dodatka br\\.': f"Dodatka br. {data.get('dodatak_broj', 'X')}",
        'Dodatkom br\\.': f"Dodatkom br. {data.get('dodatak_broj', 'X')}",
        'Dodatku br\\.': f"Dodatku br. {data.get('dodatak_broj', 'X')}",
        '\\[upisati naziv ugovora – npr\\. o nenamjenskom kreditu\\]': data.get('original_ugovor_naziv', '[NAZIV_UGOVORA]'),
        '\\[upisati naziv ugovora - npr\\. o nenamjenskom kreditu\\]': data.get('original_ugovor_naziv', '[NAZIV_UGOVORA]'),
        '\\[9910000000\\]': data.get('original_ugovor_broj_partije', '[BROJ_PARTIJE]'), # Placeholder for loan number
        '9910000000': data.get('original_ugovor_broj_partije', '[BROJ_PARTIJE]'), # Variations
        'IME I PREZIME': data.get('korisnik_ime_prezime', '[IME_PREZIME_KORISNIKA]'), # Placeholder for name (first occurrence assumed Korisnik)
        'Adresa': data.get('korisnik_adresa', '[ADRESA_KORISNIKA]'), # Placeholder for address (first occurrence assumed Korisnik)
        '___________': data.get('korisnik_oib', '[OIB_KORISNIKA]'), # Placeholder for OIB (first occurrence assumed Korisnik)

        # Specific placeholders for amounts and dates need careful handling
        # Using more specific placeholders if possible is better than generic ones.

        # Calculated/Formatted strings
        '\\[upisati datum slovima\\]': date_to_hr_words(data.get('dodatak_datum_zakljucenja', '')), # Addendum date in words
        '\\[XX\\.XXX,XX\\] \\[VALUTA\\]': f"{data.get('original_ugovor_iznos', 0.0):,.2f} {data.get('original_ugovor_valuta', 'EUR')}".replace(',', 'X').replace('.', ',').replace('X', '.'), # Original amount formatted
        '\\[upisati slovima iznos\\]': num_to_hr_words(data.get('original_ugovor_iznos', 0.0), data.get('original_ugovor_valuta', 'EUR')), # Original amount in words

        # Need to distinguish between different [DD.MM.GGGG.] and [slovima] placeholders if template uses them generically
        # Assuming specific context based on surrounding text for now.
        # Article 1 - Original Contract Date
        # Need unique placeholders in template or more context-aware replacement logic.
        # For MVP, we rely on the order or make assumptions.
        # Let's assume the FIRST [DD.MM.GGGG.] after "dana" in Article 1 is original contract date
        # Let's assume the FIRST [upisati slovima iznos] after the amount in Article 1 is original amount words

        # Article 2 - Payment, New Principal, New Annuity
        # Placeholder: u iznosu od [XX.XXX,XX] EUR (slovima: [upisati slovima iznos]) <-- Payment amount
        # Placeholder: ista sada iznosi [XX.XXX,XX] EUR (slovima: [upisati slovima iznos]) <-- New Principal
        # Placeholder: iznosi [XX.XXX,XX] EUR (slovima: [upisati slovima iznos]) <-- New Annuity

        # Placeholder for Solidarni Duznik / Jamac - need specific handling if they exist
        # Example: Identify paragraphs/tables related to them and fill if data present

        # Conditional text removal (e.g., NHB clauses, specific annuity clause)
        # Requires identifying the paragraphs/runs to remove.
        # Example: Removing the second option for Article 2, paragraph 2 (keeping the shorter term one)
    }

    # Simple text replacement in paragraphs and runs
    for para in doc.paragraphs:
        inline = para.runs
        for item in inline:
            # Simple check and replace (might need more robust regex)
            for key, value in replacements.items():
                 # Use re.sub for regex replacements in the text of each run
                 if re.search(key, item.text):
                     # Replace directly - may affect formatting if placeholder spans multiple runs
                     item.text = re.sub(key, str(value), item.text)


    # More robust replacement logic handling placeholders across multiple runs:
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\\n" # Use a unique separator

    # Perform replacements on the full text representation
    temp_full_text = full_text
    for key, value in replacements.items():
         temp_full_text = re.sub(key, str(value), temp_full_text)

    # This part is complex: Need to map the replaced text back to the docx structure
    # without losing formatting. Simple text replacement per run (above) is easier
    # but less reliable for multi-run placeholders.
    # For MVP, the simple run-based replacement might suffice if placeholders are within single runs.

    # --- Specific Replacements using paragraph analysis (more reliable) ---

    in_article_1 = False
    in_article_2 = False
    article_1_date_replaced = False
    article_1_amount_num_replaced = False
    article_1_amount_words_replaced = False
    article_2_payment_num_replaced = False
    article_2_payment_words_replaced = False
    article_2_principal_num_replaced = False
    article_2_principal_words_replaced = False
    article_2_annuity_num_replaced = False
    article_2_annuity_words_replaced = False
    article_2_clause_to_remove_found = False
    article_2_clause_to_keep_found = False

    paragraphs_to_remove = [] # Store paragraphs to be deleted later

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        # print(f"DEBUG Para {i}: {text[:100]}") # Debugging

        # Track current article
        if text.strip().startswith("Članak 1."):
            in_article_1 = True
            in_article_2 = False
            # print("DEBUG: Entered Article 1")
        elif text.strip().startswith("Članak 2."):
            in_article_1 = False
            in_article_2 = True
            # print("DEBUG: Entered Article 2")
        elif text.strip().startswith("Članak 3."):
            in_article_1 = False
            in_article_2 = False
            # print("DEBUG: Exited Article 2")


        # Generic replacements first (already done partially above)
        # Redo them here with paragraph context if needed

        # --- Article 1 specific replacements ---
        if in_article_1:
            # Original contract date
            if "dana [DD.MM.GGGG.]" in text and not article_1_date_replaced:
                para.text = para.text.replace("[DD.MM.GGGG.]", data.get('original_ugovor_datum_zakljucenja', '[DATUM_ORIG_UGOVORA]'))
                para.text = para.text.replace("[upisati datum slovima]", date_to_hr_words(data.get('original_ugovor_datum_zakljucenja', '')))
                article_1_date_replaced = True
                # print(f"DEBUG: Replaced Art 1 Date/Words in Para {i}")

            # Original contract amount
            if "[XX.XXX,XX] [VALUTA]" in text and not article_1_amount_num_replaced:
                 amount_str = f"{data.get('original_ugovor_iznos', 0.0):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                 para.text = para.text.replace("[XX.XXX,XX] [VALUTA]", f"{amount_str} {data.get('original_ugovor_valuta', 'EUR')}")
                 article_1_amount_num_replaced = True
                 # print(f"DEBUG: Replaced Art 1 Amount Num in Para {i}")

            if "(slovima: [upisati slovima iznos])" in text and article_1_amount_num_replaced and not article_1_amount_words_replaced:
                 para.text = para.text.replace("[upisati slovima iznos]", num_to_hr_words(data.get('original_ugovor_iznos', 0.0), data.get('original_ugovor_valuta', 'EUR')))
                 article_1_amount_words_replaced = True
                 # print(f"DEBUG: Replaced Art 1 Amount Words in Para {i}")

            # Remove NHB/HPB migration clauses if they exist and shouldn't be there
            # (Assuming they are separate paragraphs based on template structure)
            if "Nova hrvatska banka dioničko društvo" in text and "pripojena Banci" in text:
                 if i > 0 and doc.paragraphs[i-1].text.strip() == "UVOD": # Check if it's the intro paragraph
                      # Decide whether to keep or remove based on actual context (MVP removes if standard HPB)
                      # paragraphs_to_remove.append(para) # Schedule for removal
                      pass # Keep intro for now unless logic dictates removal
            if "nakon uvođenja EUR-a kao službene valute" in text:
                 paragraphs_to_remove.append(para) # Remove EUR conversion clause for this MVP case
                 # print(f"DEBUG: Marked Para {i} for removal (EUR clause)")
            if "nakon pripajanja pravnog prednika Banke Banci" in text:
                 paragraphs_to_remove.append(para) # Remove NHB migration clause for this MVP case
                 # print(f"DEBUG: Marked Para {i} for removal (NHB clause)")


        # --- Article 2 specific replacements ---
        if in_article_2:
             # Partial payment amount
             if "u iznosu od [XX.XXX,XX] EUR" in text and not article_2_payment_num_replaced:
                  amount_str = f"{data.get('uplata_iznos', 0.0):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                  para.text = para.text.replace("[XX.XXX,XX] EUR", f"{amount_str} EUR")
                  article_2_payment_num_replaced = True
                  if "(slovima: [upisati slovima iznos])" in text and not article_2_payment_words_replaced:
                      para.text = para.text.replace("[upisati slovima iznos]", num_to_hr_words(data.get('uplata_iznos', 0.0), "EUR"))
                      article_2_payment_words_replaced = True
                  # print(f"DEBUG: Replaced Art 2 Payment Num/Words in Para {i}")

             # New principal amount
             if "ista sada iznosi [XX.XXX,XX] EUR" in text and not article_2_principal_num_replaced:
                  amount_str = f"{data.get('nova_glavnica_iznos', 0.0):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                  para.text = para.text.replace("[XX.XXX,XX] EUR", f"{amount_str} EUR")
                  article_2_principal_num_replaced = True
                  if "(slovima: [upisati slovima iznos])" in text and not article_2_principal_words_replaced:
                       para.text = para.text.replace("[upisati slovima iznos]", num_to_hr_words(data.get('nova_glavnica_iznos', 0.0), "EUR"))
                       article_2_principal_words_replaced = True
                  # print(f"DEBUG: Replaced Art 2 Principal Num/Words in Para {i}")


             # New annuity amount / Clause handling (Keep shorter term, remove same term)
             annuity_clause_shorter_term_marker = "mijenja datum dospijeća zadnjeg anuiteta" # From template [cite: 87]
             annuity_clause_same_term_marker = "Mjesečni anuitet na preostali iznos nedospjele glavnice iznosi [XX.XXX,XX] EUR (slovima: [upisati  slovima iznos])." # From template [cite: 88] - This is too generic, need better marker if possible. Let's use presence/absence of the other marker.

             is_shorter_term_clause = annuity_clause_shorter_term_marker in text
             is_potentially_same_term_clause = text.strip().startswith("2)") and "[XX.XXX,XX] EUR" in text and not is_shorter_term_clause


             if is_shorter_term_clause and not article_2_annuity_num_replaced:
                  # This is the clause we want to keep and fill
                  amount_str = f"{data.get('novi_anuitet_iznos', 0.0):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                  para.text = para.text.replace("[XX.XXX,XX] EUR", f"{amount_str} EUR")
                  article_2_annuity_num_replaced = True
                  if "(slovima: [upisati slovima iznos])" in text and not article_2_annuity_words_replaced:
                       para.text = para.text.replace("[upisati slovima iznos]", num_to_hr_words(data.get('novi_anuitet_iznos', 0.0), "EUR"))
                       article_2_annuity_words_replaced = True
                  article_2_clause_to_keep_found = True
                  # print(f"DEBUG: Replaced Art 2 Annuity (Shorter Term) in Para {i}")

             elif is_potentially_same_term_clause and not article_2_clause_to_remove_found and not article_2_clause_to_keep_found:
                  # This looks like the clause to remove
                   paragraphs_to_remove.append(para)
                   article_2_clause_to_remove_found = True
                   # print(f"DEBUG: Marked Para {i} for removal (Same Term Annuity Clause)")

    # --- Remove paragraphs marked for deletion ---
    # Iterate backwards to avoid index issues after deletion
    # print(f"DEBUG: Removing {len(paragraphs_to_remove)} paragraphs.")
    for para in reversed(paragraphs_to_remove):
        try:
            p_element = para._element
            p_element.getparent().remove(p_element)
            # para._p = para._element = None # Mark as removed
        except Exception as e:
            print(f"WARN: Could not remove paragraph: {e}", file=sys.stderr)


    # --- Fill Party Placeholders ---
    # This needs more robust logic, e.g., finding specific labels
    # MVP: Simple replacement based on initial template structure
    party_section_found = False
    korisnik_filled = False
    solidarni_duznik_filled = False
    solidarni_jamac_filled = False

    for para in doc.paragraphs:
         text = para.text.strip()

         if "IME I PREZIME iz Adresa, OIB:" in text: # Marker for party block
             party_section_found = True

         if party_section_found:
              # Korisnik Kredita (Assume first block)
              if "IME I PREZIME" in text and not korisnik_filled:
                   para.text = para.text.replace("IME I PREZIME", data.get('korisnik_ime_prezime', '[IME_PREZIME_KORISNIKA]'))
                   if "Adresa" in para.text: # Check if Adresa is in the same para
                       para.text = para.text.replace("Adresa", data.get('korisnik_adresa', '[ADRESA_KORISNIKA]'))
                   if "___________" in para.text: # Check if OIB is in the same para
                        para.text = para.text.replace("___________", data.get('korisnik_oib', '[OIB_KORISNIKA]'))
                   korisnik_filled = True
                   # print(f"DEBUG: Filled Korisnik Kredita block")
                   continue # Move to next paragraph

              # Solidarni Dužnik (Assume second block if present)
              # Add similar logic here if solidarni_duznik data exists

              # Solidarni Jamac (Assume third/fourth block if present)
              # Add similar logic here if solidarni_jamac data exists

    # --- Fill Signature Placeholders ---
    signature_section = False
    for para in doc.paragraphs:
        if "KORISNIK KREDITA" in para.text and "ZA BANKU" in para.text:
             signature_section = True
        if signature_section:
             if "___________________" in para.text:
                 # Potentially add names above signature lines if needed
                 pass # Keep lines for actual signatures


    # Final save
    try:
        doc.save(output_path)
        print(f"Successfully generated document: {output_path}")
        return {"success": True, "output_path": output_path}
    except Exception as e:
        print(f"ERROR: Failed to save document '{output_path}'. Error: {e}", file=sys.stderr)
        return {"success": False, "error": f"Failed to save document: {e}"}

# --- Entry point for MCP execution ---
if __name__ == '__main__':
    import json
    # Expecting args as a JSON string from stdin or command line
    if len(sys.argv) > 1:
         args_json = sys.argv[1]
    else:
         args_json = sys.stdin.read()

    try:
        args = json.loads(args_json)
        result = fill_template(
            template_path=args['template_file_path'],
            output_path=args['output_file_path'],
            data=args['data']
        )
        print(json.dumps(result)) # Output result as JSON
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}), file=sys.stderr)
        sys.exit(1)

"""

# --- Agent Setup ---

# Define MCP Servers
servers = [
    MCPServerStdio(
        'npx',
        ['-y', '@pydantic/mcp-run-python', 'stdio', '--python-path', sys.executable], # Use current python env
      
    ),
    MCPServerStdio(
        'npx',
        ['-y', '@modelcontextprotocol/server-filesystem', str(BASE_WORKING_DIR)], # Serve files from BASE_WORKING_DIR
       
    )
]

# Define the AI Model
model = OpenAIModel(
    'gpt-4o', # Or another suitable model like gpt-3.5-turbo
    provider=OpenAIProvider(api_key=OPENAI_API_KEY)
)

# Create the Agent
# Note: The system prompt could be more specific about using the tools for document generation
agent = Agent(
    model=model,
    system_prompt="You are a helpful banking assistant. Your task is to fill document templates using provided data and tools. You can execute Python code and access a local filesystem.",
    mcp_servers=servers,
    # debug=True # Enable for more verbose output
)

# --- Main Execution Logic ---
async def main():
    print("=== Banking Document Assistant ===")
    print(f"Working Directory (for MCP Filesystem): {BASE_WORKING_DIR}")
    print(f"Template File: {TEMPLATE_FILENAME}")
    print(f"Output File: {OUTPUT_FILENAME}")
    print("--- Pre-extracted Data ---")
    for key, value in extracted_data.items():
        print(f"  {key}: {value}")
    print("--------------------------")
    print("Pozdrav! Što mogu učiniti za vas danas?")
    template_path_mcp = TEMPLATE_FILENAME # Path relative to BASE_WORKING_DIR
    output_path_mcp = OUTPUT_FILENAME   # Path relative to BASE_WORKING_DIR

    # Prepare the arguments for the Python code execution tool
    fill_args = FillDocxArgs(
        template_file_path=template_path_mcp,
        output_file_path=output_path_mcp,
        data=extracted_data
    )
    fill_args_json = fill_args.model_dump_json()

    # Construct the prompt for the agent to use the Python executor
    prompt = f"""
    Please execute the provided Python code to fill a DOCX template.

    Python Code to Execute:
    ```python
    {PYTHON_CODE_FOR_MCP}
    ```

    Arguments for the Python code (as JSON string):
    {fill_args_json}

    Use the 'python_executor' tool with the provided code and arguments.
    The python code will handle reading the template '{template_path_mcp}'
    and writing the output '{output_path_mcp}' via the 'filesystem' server implicitly
    (as paths are relative to its root: {BASE_WORKING_DIR}).
    Report the final result (success or failure) and the output path if successful.
    """

    print("\n[Agent Task] Asking agent to execute Python code via MCP to generate the document...")

    try:
        async with agent.run_mcp_servers():
            result = await agent.run(prompt)
            print("\n[Agent Result Raw Data]")
            print(result.output) # Agent's textual response based on tool execution

            # We expect the python_executor tool to print JSON output.
            # The agent *should* relay this, but we might need to parse the tool output directly
            # if the agent doesn't format it nicely. Pydantic AI might improve this over time.

            # Check if the output file exists using the filesystem tool (optional verification)
            # This requires the agent understanding to use the filesystem tool.
            # check_file_prompt = f"Check if the file '{output_path_mcp}' exists using the filesystem tool."
            # file_check_result = await agent.run(check_file_prompt, message_history=result.all_messages())
            # print("\n[Agent File Check]")
            # print(file_check_result.data)

            # Manual check based on expected output path
            output_file_full_path = BASE_WORKING_DIR / OUTPUT_FILENAME
            if output_file_full_path.exists():
                 print(f"\n[Verification] SUCCESS: Output file '{output_file_full_path}' seems to have been created.")
            else:
                 print(f"\n[Verification] WARNING: Output file '{output_file_full_path}' was NOT found. Check agent logs and MCP server output for errors.")


    except Exception as e:
        print(f"\nError during agent execution: {e}")
        print("Please check:")
        print("1. Your OpenAI API key in the .env file.")
        print("2. The BASE_WORKING_DIR path is correct and accessible.")
        print("3. The template file exists at the expected location.")
        print("4. `npx` is installed and accessible.")
        print("5. The required Python libraries (`python-docx`, `num2words`) are installed in the environment accessible to '@pydantic/mcp-run-python'.")
        print("6. MCP server logs (if any are produced in the console).")

if __name__ == '__main__':
    # Ensure the working directory exists
    if not BASE_WORKING_DIR.exists():
         print(f"ERROR: Base working directory '{BASE_WORKING_DIR}' does not exist. Please create it or update the path.")
         sys.exit(1)
    # Ensure the template file exists
    if not (BASE_WORKING_DIR / TEMPLATE_FILENAME).exists():
         print(f"ERROR: Template file '{TEMPLATE_FILENAME}' not found in '{BASE_WORKING_DIR}'.")
         sys.exit(1)

    asyncio.run(main())