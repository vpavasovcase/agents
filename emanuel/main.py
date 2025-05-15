#!/usr/bin/env python3
"""
CLI‑based AI agent system for filling out a loan‑agreement template.

Updated to demonstrate **programmatic agent hand‑off** à la the example you
provided (flight→seat).  Each LLM‑driven stage is now a distinct
`pydantic_ai.Agent` with its own `output_type`, while non‑LLM stages remain as
plain functions.  The control‑flow decides which agent / function to call next
based on intermediate results.

▸ **Environment & secrets** via `python‑dotenv` (`.env` with `GROQ_API_KEY=…`).
▸ Uses two MCP servers (filesystem + Office‑Word) for safe document access.
▸ Works with any set of `{{placeholders}}`; nothing is hard‑coded.
"""

import os
import re
import asyncio
from pathlib import Path
from typing import List, Dict, Tuple

from dotenv import load_dotenv

import logfire
from pydantic import BaseModel
from pydantic_ai import Agent, RunContext
from pydantic_ai.models.groq import GroqModel
from pydantic_ai.providers.groq import GroqProvider
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai.messages import ModelMessage
from pydantic_ai.usage import Usage, UsageLimits

# OCR / document libs ----------------------------------------------------------
from pdf2image import convert_from_path
import pytesseract
import fitz  # PyMuPDF
import docx  # python‑docx

# ─── logging & instrumentation ────────────────────────────────────────────────
load_dotenv()
logfire.configure()
Agent.instrument_all()

# ─── LLM model & MCP servers --------------------------------------------------
llm_model = GroqModel(
    "meta‑llama/llama‑4‑maverick‑17b‑128e‑instruct",
    provider=GroqProvider(api_key=os.getenv("GROQ_API_KEY", "")),
)

DOCS_BASE_PATH = Path("docs").resolve().as_posix()

mcp_servers = [
    MCPServerStdio(
        "npx", ["-y", "@modelcontextprotocol/server-filesystem", DOCS_BASE_PATH]
    ),
    MCPServerStdio(
        "uvx", ["--from", "office-word-mcp-server", "word_mcp_server"]
    ),
]

# ─── pydantic output schemas for agents ---------------------------------------
class RequiredFields(BaseModel):
    required_fields: List[str]


class SearchOutput(BaseModel):
    found_data: Dict[str, str]
    missing_data: List[str]


# ─── AGENT #1 ─ AnalysisAgent  (extract placeholders) -------------------------
analysis_agent = Agent[None, RequiredFields](
    model=llm_model,
    output_type=RequiredFields,  # type: ignore
    system_prompt=(
        "You are the Analysis Agent. Given a list of placeholder names from a "
        "Word template and comments extracted from a PDF version, decide which "
        "placeholders must be filled out by downstream agents. Return JSON "
        "{\"required_fields\": [...]}."
    ),
    mcp_servers=mcp_servers,
)

# ─── AGENT #2 ─ SearchAgent  (map OCR→values) --------------------------------
search_agent = Agent[None, SearchOutput](
    model=llm_model,
    output_type=SearchOutput,  # type: ignore
    system_prompt=(
        "You are the Search Agent. For each required field, extract the best "
        "matching value from the provided document text. If not found or "
        "ambiguous, list the field under 'missing_data'. Return JSON "
        "{\"found_data\": {...}, \"missing_data\": [...]}."
    ),
    mcp_servers=mcp_servers,
)

# ─── NON‑LLM utility funcs ----------------------------------------------------

def ocr_all_pdfs(credit: str) -> str:
    """OCR every PDF in docs/sources/<credit>/ and concatenate."""
    src = Path(f"docs/sources/{credit}")
    if not src.exists():
        print(f"[error] No source directory {src}")
        return ""
    chunks = []
    for fp in src.iterdir():
        if fp.suffix.lower() == ".pdf":
            try:
                for img in convert_from_path(fp):
                    chunks.append(pytesseract.image_to_string(img))
            except Exception as exc:
                print(f"[warn] OCR failed on {fp}: {exc}")
        elif fp.suffix.lower() == ".txt":
            chunks.append(fp.read_text(encoding="utf-8"))
    combined = "\n".join(chunks)
    logfire.info("read_agent", credit=credit, chars=len(combined))
    return combined


def ask_user(found: Dict[str, str], missing: List[str]) -> None:
    for field in missing:
        ans = input(f"Enter value for '{field}': ").strip()
        if ans:
            found[field] = ans


def placeholders_from_docx(docx_path: str) -> List[str]:
    doc = docx.Document(docx_path)
    ph = {m for p in doc.paragraphs for m in re.findall(r"\{\{(\w+)\}\}", p.text)}
    for tbl in doc.tables:
        for cell in tbl._cells:
            ph.update(re.findall(r"\{\{(\w+)\}\}", cell.text))
    return sorted(ph)


def comments_from_pdf(pdf_path: str) -> List[str]:
    out: List[str] = []
    pdf = fitz.open(pdf_path)
    for page in pdf:
        annot = page.first_annot
        while annot:
            if (txt := annot.get_text()):
                out.append(txt.strip())
            annot = annot.next
    pdf.close()
    return out


def fill_docx(template: str, out_path: str, data: Dict[str, str]) -> None:
    doc = docx.Document(template)
    def sub(run):
        for k, v in data.items():
            run.text = run.text.replace(f"{{{{{k}}}}}", str(v))
    for p in doc.paragraphs:
        for r in p.runs:
            sub(r)
    for tbl in doc.tables:
        for cell in tbl._cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    sub(r)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    logfire.info("fill_agent", path=out_path)


def unresolved_placeholders(path: str) -> List[str]:
    doc = docx.Document(path)
    left = [p.text for p in doc.paragraphs if "{{" in p.text and "}}" in p.text]
    for tbl in doc.tables:
        for cell in tbl._cells:
            if "{{" in cell.text and "}}" in cell.text:
                left.append(cell.text)
    return left

# ─── Main orchestrator (programmatic hand‑off) ───────────────────────────────
async def main():
    usage = Usage()  # optional; can feed into agent.run for limits / accounting
    limits = UsageLimits(request_limit=25)

    while True:
        credit = input("Credit number (or 'exit'): ").strip()
        if credit.lower() in {"exit", "quit", ""}:
            break

        # Stage 1: Read documents → OCR text
        doc_text = ocr_all_pdfs(credit)
        if not doc_text:
            continue

        # Stage 2: AnalysisAgent decides which placeholders matter
        placeholders = placeholders_from_docx("docs/template.docx")
        comments = comments_from_pdf("docs/template.pdf")
        analysis_prompt = (
            f"PLACEHOLDERS: {placeholders}\nPDF_COMMENTS: {comments}"
        )
        res_analysis = await analysis_agent.run(
            analysis_prompt,
            usage=usage,
            usage_limits=limits,
        )
        required = res_analysis.output.required_fields

        # Stage 3: SearchAgent maps OCR text → values / missing list
        search_prompt = (
            f"REQUIRED: {required}\nTEXT: {doc_text[:1600]}"
        )
        res_search = await search_agent.run(
            search_prompt,
            usage=usage,
            usage_limits=limits,
        )
        found, missing = res_search.output.found_data, res_search.output.missing_data

        # Stage 4: Ask user for any remaining data
        if missing:
            ask_user(found, missing)

        # Ensure everything filled
        still_missing = [f for f in required if not found.get(f)]
        if still_missing:
            print("[error] Still missing values for:", still_missing)
            continue

        # Stage 5: Fill DOCX & validate
        out_docx = f"docs/completed/{credit}.docx"
        fill_docx("docs/template.docx", out_docx, found)

        leftovers = unresolved_placeholders(out_docx)
        if leftovers:
            print("❌  Unresolved placeholders remain:", leftovers)
        else:
            print("✅  Completed:", out_docx)


if __name__ == "__main__":
    asyncio.run(main())
