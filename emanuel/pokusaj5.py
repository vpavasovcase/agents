import os
import asyncio
import docx
import io
from pydantic_ai import Agent
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.providers.openai import OpenAIProvider
from dotenv import load_dotenv

# Load environment variables
load_dotenv(override=True)

# Configure available MCP servers
servers = [
    MCPServerStdio('npx', ['-y', '@pydantic/mcp-run-python', 'stdio']),
    MCPServerStdio('npx', ['-y', '@modelcontextprotocol/server-filesystem', '.'])
]

# Configure the model - using GPT-4o for strong document understanding
model = OpenAIModel(
    'gpt-4o',
    provider=OpenAIProvider(api_key=os.getenv('OPENAI_API_KEY'))
)

# Create the agent with specific system prompt for document processing
agent = Agent(
    model=model,
    system_prompt="""You are a specialized document processing assistant for a bank. 
    Your task is to analyze template documents, identify fields that need to be filled,
    extract relevant information from other provided documents, and then fill the template
    with the correct data. Be precise and accurate with financial information.
    
    When you don't have enough information or are uncertain about any field, ask the user
    for clarification instead of guessing. Banking documents require 100% accuracy.
    """,
    mcp_servers=servers
)

async def read_file_content(file_path):
    """Read the content of a file using the filesystem MCP server"""
    try:
        prompt = f"""
        Read the contents of the file at '{file_path}' and return the full text content.
        """
        result = await agent.run(prompt)
        return result.output
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return None

async def list_documents(directory):
    """List available documents in the specified directory"""
    try:
        prompt = f"""
        List all files in the directory '{directory}' with their full paths.
        Only include .txt, .docx, .pdf, and other document files.
        """
        result = await agent.run(prompt)
        print(f"Available documents in {directory}:")
        print(result.output)
        return result.output
    except Exception as e:
        print(f"Error listing documents: {e}")
        return None

async def analyze_template_content(template_content):
    """Analyze template content to identify fields that need to be filled"""
    try:
        prompt = f"""
        Analyze the following template document content:
        
        {template_content}
        
        Identify all fields that need to be filled (marked with [placeholders] or {{.mark}} tags or other indicators). 
        Create a structured list of all required information we need to collect from other documents.
        """
        result = await agent.run(prompt)
        print("\nTemplate Analysis:")
        print(result.output)
        return result.output
    except Exception as e:
        print(f"Error analyzing template: {e}")
        return None

async def extract_data_from_document_contents(document_contents, required_fields):
    """Extract required information from document contents"""
    try:
        # Combine all document contents with clear separation
        combined_content = ""
        for i, (doc_path, content) in enumerate(document_contents.items()):
            combined_content += f"\n\n--- DOCUMENT {i+1}: {doc_path} ---\n\n{content}"
        
        prompt = f"""
        I need to extract specific information from these documents to fill a bank template:
        
        {combined_content}
        
        Please extract the following information:
        {required_fields}
        
        For each field, provide:
        1. The exact value found
        2. The source document where it was found
        3. Your confidence level (high/medium/low)
        
        If some information cannot be found, mark it as "Not Found" and I'll ask the user for it.
        """
        result = await agent.run(prompt)
        print("\nExtracted Data:")
        print(result.output)
        return result.output
    except Exception as e:
        print(f"Error extracting data: {e}")
        return None

async def ask_for_missing_data(missing_fields):
    """Ask the user for any missing information"""
    user_provided_data = {}
    
    print("\nSome information is missing or needs verification:")
    for field in missing_fields:
        if field.strip():  # Only process non-empty fields
            value = input(f"{field}: ")
            user_provided_data[field] = value
    
    return user_provided_data

async def generate_filled_document_content(template_content, extracted_data):
    """Generate content for a new document with extracted data"""
    try:
        prompt = f"""
        Using the following template:
        
        {template_content}
        
        And the following extracted data:
        
        {extracted_data}
        
        Generate the complete content for a new document with all fields filled in.
        Return only the filled document content, ready to be saved.
        
        For any fields where data is missing or has low confidence, leave the placeholder
        or mark it with [NEEDS_VERIFICATION].
        """
        result = await agent.run(prompt)
        return result.output
    except Exception as e:
        print(f"Error generating filled document: {e}")
        return None

async def write_file_content(file_path, content):
    """Write content to a file using the filesystem MCP server"""
    try:
        prompt = f"""
        Write the following content to the file '{file_path}':
        
        {content}
        """
        result = await agent.run(prompt)
        print(f"\nFile written to: {file_path}")
        return result.output
    except Exception as e:
        print(f"Error writing file: {e}")
        return None

async def generate_docx_from_text(template_path, filled_content, output_path):
    """Use Python's python-docx to generate a document from template and filled content"""
    try:
        prompt = f"""
        Execute the following Python code using the @pydantic/mcp-run-python server:
        
        ```python
        import docx
        from docx import Document
        import re
        import io
        
        # Try to open the template document as a base
        try:
            doc = Document('{template_path}')
            # Create a new document based on template structure
            # We'll replace content in each paragraph
        except Exception as e:
            # If we can't open template, create a new document
            doc = Document()
            # Add a title
            doc.add_heading('Filled Bank Document', 0)
            
        # Clear existing content from template
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                for run in paragraph.runs:
                    run.text = ''
            
            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    run.text = ''
        
        # Clear main document content
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.text = ''
        
        # Add our filled content as paragraphs
        content_lines = '''{filled_content}'''.split('\\n')
        
        # Process each line of our content
        for line in content_lines:
            if line.strip():
                para = doc.add_paragraph()
                para.add_run(line)
        
        # Save the document
        doc.save('{output_path}')
        
        print(f"Document successfully created at: {output_path}")
        ```
        """
        result = await agent.run(prompt)
        print("\nDocument Creation Result:")
        print(result.output)
        return True
    except Exception as e:
        print(f"Error generating document: {e}")
        return False

async def main():
    print("=== Bank Document Template Filling Agent ===")
    print("This agent will fill a template document with data from other documents.")
    
    # Define working directory
    working_dir = "/app/emanuel"
    
    async with agent.run_mcp_servers():
        while True:
            print("\n===== Main Menu =====")
            print("1. List available documents")
            print("2. Process documents and fill template")
            print("3. Exit")
            
            choice = input("\nChoice (1-3): ")
            
            if choice == "1":
                await list_documents(working_dir)
            
            elif choice == "2":
                # Get document paths
                template_path = input("Enter the path to the template document: ")
                output_path = input("Enter the path for the output document: ")
                
                # Get other document paths
                other_docs = []
                while True:
                    doc_path = input("Enter path to a source document (or 'done' to finish): ")
                    if doc_path.lower() == 'done':
                        break
                    other_docs.append(doc_path)
                
                # Read template content
                print("\nReading template document...")
                template_content = await read_file_content(template_path)
                
                # Analyze template to identify required fields
                print("\nAnalyzing template for required fields...")
                required_fields = await analyze_template_content(template_content)
                
                # Read content of other documents
                print("\nReading source documents...")
                document_contents = {}
                for doc_path in other_docs:
                    content = await read_file_content(doc_path)
                    if content:
                        document_contents[doc_path] = content
                
                # Extract data from documents
                print("\nExtracting data from documents...")
                extracted_data = await extract_data_from_document_contents(document_contents, required_fields)
                
                # Identify missing or low-confidence data
                print("\nIdentifying missing information...")
                missing_fields_prompt = f"""
                Based on the extracted data:
                {extracted_data}
                
                Please identify any missing fields or fields with low confidence that we should ask the user about.
                Return ONLY the field names, one per line. Don't include any explanations or other text.
                """
                missing_fields_result = await agent.run(missing_fields_prompt)
                missing_fields = missing_fields_result.data.strip().split('\n')
                
                # Ask user for missing data
                user_data = await ask_for_missing_data(missing_fields)
                
                # Format user data
                user_data_formatted = "\n".join([f"{field}: {value}" for field, value in user_data.items()])
                
                # Combine extracted and user-provided data
                complete_data = f"""
                Extracted data:
                {extracted_data}
                
                User-provided data:
                {user_data_formatted}
                """
                
                # Generate filled document content
                print("\nGenerating filled document content...")
                filled_content = await generate_filled_document_content(template_content, complete_data)
                
                # Generate the document
                print("\nCreating final document...")
                success = await generate_docx_from_text(template_path, filled_content, output_path)
                
                if success:
                    print(f"\nDocument successfully created at: {output_path}")
                else:
                    # Try simple text output as fallback
                    print("\nFalling back to text file output...")
                    await write_file_content(output_path + ".txt", filled_content)
                    print(f"\nText version saved at: {output_path}.txt")
            
            elif choice == "3":
                print("Exiting. Goodbye!")
                break
            
            else:
                print("Invalid choice. Please try again.")

if __name__ == "__main__":
    asyncio.run(main())