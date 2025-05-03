import os
import asyncio
import json
import hashlib
import time
import random
import docx
import io
from pydantic_ai import Agent
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.providers.openai import OpenAIProvider
from pydantic_ai.exceptions import ModelHTTPError
from dotenv import load_dotenv

# Load environment variables
load_dotenv(override=True)

# Configure available MCP servers
servers = [
    MCPServerStdio('npx', ['-y', '@pydantic/mcp-run-python', 'stdio']),
    MCPServerStdio('npx', ['-y', '@modelcontextprotocol/server-filesystem', '.'])
]

# Define cache directory
CACHE_DIR = "./.doc_cache"
os.makedirs(CACHE_DIR, exist_ok=True)

# Configure the models - using different models for different tasks
file_model = OpenAIModel(
    'gpt-3.5-turbo',  # Use cheaper model for simple tasks
    provider=OpenAIProvider(api_key=os.getenv('OPENAI_API_KEY'))
)

analysis_model = OpenAIModel(
    'gpt-4o',  # Use more capable model for complex analysis
    provider=OpenAIProvider(api_key=os.getenv('OPENAI_API_KEY'))
)

# Create agents for different tasks
file_agent = Agent(
    model=file_model,
    system_prompt="""You are a document processing assistant. 
    Your task is to accurately read and process document content without interpretation.""",
    mcp_servers=servers
)

analysis_agent = Agent(
    model=analysis_model,
    system_prompt="""You are a specialized document processing assistant for a bank. 
    Your task is to analyze template documents, identify fields that need to be filled,
    extract relevant information from other provided documents, and then fill the template
    with the correct data. Be precise and accurate with financial information.
    
    When you don't have enough information or are uncertain about any field, ask the user
    for clarification instead of guessing. Banking documents require 100% accuracy.
    """,
    mcp_servers=servers
)

# Cache helper functions
def get_cache_key(action, input_text):
    """Generate a cache key based on action and input"""
    key = f"{action}_{hashlib.md5(input_text.encode()).hexdigest()}"
    return key

def get_cached_result(cache_key):
    """Retrieve cached result if it exists"""
    cache_file = os.path.join(CACHE_DIR, f"{cache_key}.json")
    if os.path.exists(cache_file):
        with open(cache_file, 'r') as f:
            try:
                cached_data = json.load(f)
                print(f"Using cached result for {cache_key}")
                return cached_data
            except json.JSONDecodeError:
                print(f"Cache file corrupted: {cache_file}")
                return None
    return None

def save_to_cache(cache_key, result):
    """Save result to cache"""
    cache_file = os.path.join(CACHE_DIR, f"{cache_key}.json")
    with open(cache_file, 'w') as f:
        json.dump(result, f)
    print(f"Saved result to cache: {cache_key}")

async def retry_agent_run(agent, prompt, use_cache=True, max_retries=5):
    """Run agent with retry logic and caching"""
    if use_cache:
        cache_key = get_cache_key("run", prompt[:100])
        cached_result = get_cached_result(cache_key)
        if cached_result:
            # Create a result-like object
            class CachedResult:
                def __init__(self, data):
                    self.output = data.get("output", "")
                    self.data = data.get("data", "")
            return CachedResult(cached_result)
    
    # Add jitter to avoid synchronized requests
    await asyncio.sleep(random.uniform(0.5, 2.0))
    
    retry_count = 0
    while retry_count < max_retries:
        try:
            result = await agent.run(prompt)
            
            if use_cache:
                # Save to cache
                result_data = {"output": result.output, "data": result.data}
                save_to_cache(cache_key, result_data)
            
            return result
        except Exception as e:
            retry_count += 1
            if "rate_limit" in str(e).lower():
                print(f"Rate limit hit. Waiting before retry {retry_count}/{max_retries}...")
                # Exponential backoff
                wait_time = (2 ** retry_count) + random.uniform(0, 1)
                print(f"Waiting {wait_time:.2f} seconds...")
                await asyncio.sleep(wait_time)
            else:
                print(f"Error: {e}. Retry {retry_count}/{max_retries}")
                await asyncio.sleep(1)
            
            # If we've reached max retries, raise the exception
            if retry_count >= max_retries:
                print(f"Maximum retries ({max_retries}) reached. Giving up.")
                raise e

async def read_file_content(file_path):
    """Read the content of a file using the filesystem MCP server"""
    try:
        cache_key = get_cache_key("read_file", file_path)
        cached_result = get_cached_result(cache_key)
        if cached_result:
            return cached_result.get("content", "")
        
        prompt = f"""
        Read the contents of the file at '{file_path}' and return the full text content.
        """
        result = await retry_agent_run(file_agent, prompt)
        
        # Cache the result
        save_to_cache(cache_key, {"content": result.output})
        
        return result.output
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return None

async def list_documents(directory):
    """List available documents in the specified directory"""
    try:
        prompt = f"""
        List all files in the directory '{directory}' with their full paths.
        Only include .txt, .docx, .pdf, and other document files.
        """
        result = await retry_agent_run(file_agent, prompt)
        print(f"Available documents in {directory}:")
        print(result.output)
        return result.output
    except Exception as e:
        print(f"Error listing documents: {e}")
        return None

async def analyze_template_content(template_content):
    """Analyze template content to identify fields that need to be filled"""
    try:
        cache_key = get_cache_key("analyze_template", template_content[:200])
        cached_result = get_cached_result(cache_key)
        if cached_result:
            return cached_result.get("analysis", "")
        
        prompt = f"""
        Analyze the following template document content:
        
        {template_content}
        
        Identify all fields that need to be filled (marked with [placeholders] or {{.mark}} tags or other indicators). 
        Create a structured list of all required information we need to collect from other documents.
        """
        result = await retry_agent_run(analysis_agent, prompt)
        print("\nTemplate Analysis:")
        print(result.output)
        
        # Cache the result
        save_to_cache(cache_key, {"analysis": result.output})
        
        return result.output
    except Exception as e:
        print(f"Error analyzing template: {e}")
        return None

async def extract_data_from_document_contents(document_contents, required_fields):
    """Extract required information from document contents"""
    try:
        # Create a deterministic key for caching
        combined_paths = "_".join(sorted(document_contents.keys()))
        cache_key = get_cache_key(f"extract_data_{combined_paths}", required_fields[:100])
        cached_result = get_cached_result(cache_key)
        if cached_result:
            return cached_result.get("extracted_data", "")
        
        # Process documents in batches to avoid token limits
        all_extracted_data = []
        
        # Group documents into batches of 2
        doc_items = list(document_contents.items())
        batches = [doc_items[i:i+2] for i in range(0, len(doc_items), 2)]
        
        for batch_index, batch in enumerate(batches):
            # Combine the batch documents with clear separation
            batch_content = ""
            for i, (doc_path, content) in enumerate(batch):
                batch_content += f"\n\n--- DOCUMENT {batch_index*2 + i+1}: {doc_path} ---\n\n{content}"
            
            prompt = f"""
            I need to extract specific information from these documents to fill a bank template:
            
            {batch_content}
            
            Please extract the following information:
            {required_fields}
            
            For each field, provide:
            1. The exact value found
            2. The source document where it was found
            3. Your confidence level (high/medium/low)
            
            If some information cannot be found, mark it as "Not Found" and I'll ask the user for it.
            """
            
            print(f"\nProcessing batch {batch_index+1}/{len(batches)}...")
            result = await retry_agent_run(analysis_agent, prompt)
            all_extracted_data.append(result.output)
            
            # Add delay between batches
            if batch_index < len(batches) - 1:
                await asyncio.sleep(2)
        
        # Combine all extracted data
        combined_data = "\n\n".join(all_extracted_data)
        
        # Now consolidate the data from all batches
        consolidation_prompt = f"""
        I have extracted data from multiple document batches. Please consolidate this information
        into a single, coherent set of extracted data, removing any duplicates and resolving
        any conflicts (choose the higher confidence data when there are conflicts):
        
        {combined_data}
        
        Provide a clean, consolidated list of all extracted fields with their values, sources, and confidence levels.
        """
        
        consolidated_result = await retry_agent_run(analysis_agent, consolidation_prompt)
        print("\nExtracted and Consolidated Data:")
        print(consolidated_result.output)
        
        # Cache the consolidated result
        save_to_cache(cache_key, {"extracted_data": consolidated_result.output})
        
        return consolidated_result.output
    except Exception as e:
        print(f"Error extracting data: {e}")
        return None

async def ask_for_missing_data(extracted_data):
    """Ask the user for any missing information"""
    try:
        # Identify missing fields
        prompt = f"""
        Based on the extracted data:
        {extracted_data}
        
        Please identify any missing fields or fields with low confidence that we should ask the user about.
        Return ONLY the field names, one per line. Don't include any explanations or other text.
        If there are no missing fields, respond with "NONE".
        """
        result = await retry_agent_run(analysis_agent, prompt)
        missing_fields = result.data.strip().split('\n')
        
        # Filter out empty fields and check for "NONE"
        missing_fields = [field for field in missing_fields if field.strip()]
        if len(missing_fields) == 1 and missing_fields[0].upper() == "NONE":
            print("\nNo missing information identified.")
            return {}
        
        user_provided_data = {}
        
        print("\nSome information is missing or needs verification:")
        for field in missing_fields:
            if field.strip():  # Only process non-empty fields
                value = input(f"{field}: ")
                user_provided_data[field] = value
        
        return user_provided_data
    except Exception as e:
        print(f"Error identifying missing fields: {e}")
        print("\nPlease manually provide any missing information:")
        field = input("Field name (or 'done' to finish): ")
        user_provided_data = {}
        
        while field.lower() != 'done':
            value = input(f"Value for {field}: ")
            user_provided_data[field] = value
            field = input("Field name (or 'done' to finish): ")
        
        return user_provided_data

async def generate_filled_document_content(template_content, extracted_data):
    """Generate content for a new document with extracted data"""
    try:
        prompt = f"""
        Using the following template:
        
        {template_content}
        
        And the following extracted data:
        
        {extracted_data}
        
        Generate the complete content for a new document with all fields filled in.
        Return only the filled document content, ready to be saved.
        
        For any fields where data is missing or has low confidence, leave the placeholder
        or mark it with [NEEDS_VERIFICATION].
        """
        result = await retry_agent_run(analysis_agent, prompt)
        return result.output
    except Exception as e:
        print(f"Error generating filled document: {e}")
        return None

async def write_file_content(file_path, content):
    """Write content to a file using the filesystem MCP server"""
    try:
        prompt = f"""
        Write the following content to the file '{file_path}':
        
        {content}
        """
        result = await retry_agent_run(file_agent, prompt, use_cache=False)
        print(f"\nFile written to: {file_path}")
        return result.output
    except Exception as e:
        print(f"Error writing file: {e}")
        return None

async def generate_docx_from_text(template_path, filled_content, output_path):
    """Use Python's python-docx to generate a document from template and filled content"""
    try:
        # Split content into smaller chunks to avoid token limits
        content_parts = [filled_content[i:i+4000] for i in range(0, len(filled_content), 4000)]
        
        for i, part in enumerate(content_parts):
            part_prompt = f"""
            Execute the following Python code using the @pydantic/mcp-run-python server:
            
            ```python
            import docx
            from docx import Document
            import re
            import io
            import os
            
            # Define output path
            output_path = '{output_path}'
            is_first_part = {i == 0}
            
            # For first part, create or open the document
            if is_first_part:
                try:
                    # Try to open the template document as a base
                    doc = Document('{template_path}')
                    # We'll keep the template structure but replace content
                except Exception as e:
                    # If we can't open template, create a new document
                    doc = Document()
                    doc.add_heading('Filled Bank Document', 0)
                
                # Clear existing content from template
                for section in doc.sections:
                    for paragraph in section.header.paragraphs:
                        for run in paragraph.runs:
                            run.text = ''
                    
                    for paragraph in section.footer.paragraphs:
                        for run in paragraph.runs:
                            run.text = ''
                
                # Clear main document content
                for paragraph in list(doc.paragraphs):
                    for run in paragraph.runs:
                        run.text = ''
            else:
                # For subsequent parts, open the existing document
                doc = Document(output_path)
            
            # Add our filled content as paragraphs
            content_lines = '''{part}'''.split('\\n')
            
            # Process each line of our content
            for line in content_lines:
                if line.strip():
                    para = doc.add_paragraph()
                    para.add_run(line)
            
            # Save the document
            doc.save(output_path)
            
            print(f"Document part {i+1} successfully added to: {output_path}")
            ```
            """
            
            print(f"\nProcessing document part {i+1}/{len(content_parts)}...")
            part_result = await retry_agent_run(file_agent, part_prompt, use_cache=False)
            print(part_result.output)
            
            # Add small delay between parts
            if i < len(content_parts) - 1:
                await asyncio.sleep(1)
        
        return True
    except Exception as e:
        print(f"Error generating document: {e}")
        return False

async def main():
    print("=== Bank Document Template Filling Agent ===")
    print("This agent will fill a template document with data from other documents.")
    
    # Define working directory
    working_dir = "/app/emanuel"
    
    async with file_agent.run_mcp_servers(), analysis_agent.run_mcp_servers():
        while True:
            print("\n===== Main Menu =====")
            print("1. List available documents")
            print("2. Process documents and fill template")
            print("3. Clear cache")
            print("4. Exit")
            
            choice = input("\nChoice (1-4): ")
            
            if choice == "1":
                await list_documents(working_dir)
            
            elif choice == "2":
                # Get document paths
                template_path = input("Enter the path to the template document: ")
                output_path = input("Enter the path for the output document: ")
                
                # Add .docx extension if not provided
                if not output_path.lower().endswith(('.docx', '.doc')):
                    output_path += ".docx"
                
                # Get other document paths
                other_docs = []
                while True:
                    doc_path = input("Enter path to a source document (or 'done' to finish): ")
                    if doc_path.lower() == 'done':
                        break
                    other_docs.append(doc_path)
                
                # Read template content
                print("\nReading template document...")
                template_content = await read_file_content(template_path)
                
                if not template_content:
                    print("Failed to read template document. Please check the path and try again.")
                    continue
                
                # Analyze template to identify required fields
                print("\nAnalyzing template for required fields...")
                required_fields = await analyze_template_content(template_content)
                
                if not required_fields:
                    print("Failed to analyze template. Please try again.")
                    continue
                
                # Read content of other documents
                print("\nReading source documents...")
                document_contents = {}
                for doc_path in other_docs:
                    print(f"Reading {doc_path}...")
                    content = await read_file_content(doc_path)
                    if content:
                        document_contents[doc_path] = content
                        print(f"Successfully read {doc_path}")
                    else:
                        print(f"Failed to read {doc_path}")
                
                if not document_contents:
                    print("Failed to read any source documents. Please check the paths and try again.")
                    continue
                
                # Extract data from documents
                print("\nExtracting data from documents...")
                extracted_data = await extract_data_from_document_contents(document_contents, required_fields)
                
                if not extracted_data:
                    print("Failed to extract data. Please try again.")
                    continue
                
                # Ask user for missing data
                user_data = await ask_for_missing_data(extracted_data)
                
                # Format user data
                user_data_formatted = "\n".join([f"{field}: {value}" for field, value in user_data.items()])
                
                # Combine extracted and user-provided data
                complete_data = f"""
                Extracted data:
                {extracted_data}
                
                User-provided data:
                {user_data_formatted}
                """
                
                # Generate filled document content
                print("\nGenerating filled document content...")
                filled_content = await generate_filled_document_content(template_content, complete_data)
                
                if not filled_content:
                    print("Failed to generate filled document content. Please try again.")
                    continue
                
                # Generate the document
                print("\nCreating final document...")
                success = await generate_docx_from_text(template_path, filled_content, output_path)
                
                if success:
                    print(f"\nDocument successfully created at: {output_path}")
                else:
                    # Try simple text output as fallback
                    print("\nFalling back to text file output...")
                    text_output_path = output_path + ".txt"
                    await write_file_content(text_output_path, filled_content)
                    print(f"\nText version saved at: {text_output_path}")
            
            elif choice == "3":
                # Clear cache
                confirm = input("Are you sure you want to clear the cache? (y/n): ")
                if confirm.lower() == 'y':
                    for cache_file in os.listdir(CACHE_DIR):
                        if cache_file.endswith('.json'):
                            os.remove(os.path.join(CACHE_DIR, cache_file))
                    print("Cache cleared successfully.")
            
            elif choice == "4":
                print("Exiting. Goodbye!")
                break
            
            else:
                print("Invalid choice. Please try again.")

if __name__ == "__main__":
    asyncio.run(main())