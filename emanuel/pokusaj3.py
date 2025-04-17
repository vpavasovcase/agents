# filename: emanuel/pokusaj2.py
# (or whatever you choose to name the file)

import asyncio
import os
import re
import sys
import json
from pydantic_ai.mcp import MCPServerStdio
from datetime import datetime
from pathlib import Path

servers = [
    MCPServerStdio('npx', ['-y', '@pydantic/mcp-run-python', 'stdio']),
    MCPServerStdio('npx', [
              "-y",
              "@modelcontextprotocol/server-filesystem",
              "/app"
            ]),
]

# Attempt to import dependencies, provide guidance if missing
try:
    from dotenv import load_dotenv
except ImportError:
    print("ERROR: 'python-dotenv' library not found. Please install it (`pip install python-dotenv`).")
    sys.exit(1)
try:
    from num2words import num2words
except ImportError:
    print("ERROR: 'num2words' library not found. Please install it (`pip install num2words`).")
    sys.exit(1)
try:
    from pydantic import BaseModel, Field
except ImportError:
    print("ERROR: 'pydantic' library not found. Please install it (`pip install pydantic`).")
    sys.exit(1)
try:
    # Pydantic AI imports
    from pydantic_ai import Agent
    from pydantic_ai.mcp import MCPServerStdio
    from pydantic_ai.models.openai import OpenAIModel
    from pydantic_ai.providers.openai import OpenAIProvider
except ImportError:
     print("ERROR: 'pydantic-ai' library not found. Please install it (`pip install \"pydantic-ai[openai]\"`).")
     sys.exit(1)

# Docx generation imports (will be used inside the MCP Python execution)
# Ensure these are importable in the environment where mcp-run-python runs
try:
    from docx import Document
    from docx.shared import Pt
    # from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Uncomment if needed
except ImportError:
    # These are not strictly needed in the main script's environment,
    # but helps with clarity / local testing if desired.
    print("INFO: 'python-docx' library not found in the main environment. "
          "Ensure it's available for the '@pydantic/mcp-run-python' server.")
    Document = None # type: ignore

# Load environment variables from .env file
load_dotenv(override=True)

# --- Configuration ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    print("ERROR: OPENAI_API_KEY not found in environment variables. Please set it in your .env file.")
    sys.exit(1)


# *** IMPORTANT: Update this path to the directory where your files are located ***
# This path will be used by the MCP filesystem server
# Use an environment variable or make it easily configurable
DEFAULT_WORKING_DIR = "/app/emanuel" # Example for your Docker setup, adjust if needed
BASE_WORKING_DIR_STR = os.getenv("AGENT_WORKING_DIR", DEFAULT_WORKING_DIR)
BASE_WORKING_DIR = Path(BASE_WORKING_DIR_STR)


TEMPLATE_FILENAME = "Predložak Dodatka  ugovoru _ UVEĆANA UPLATA_DJELOMIČNA OTPLATA KREDITA _bez izmjene roka _ kraći rok.docx"
OUTPUT_FILENAME = "POPUNJEN_Dodatak_ugovoru_Barbara_Stazic.docx"

# Source documents (relative to BASE_WORKING_DIR) - We won't read them directly here,
# data is pre-extracted for MVP, but listing for completeness.
# SOURCE_PDF_STATUS = "DocumentServlet.pdf"
# SOURCE_PDF_CONTRACT = "FileDownloadServlet.pdf"
# SOURCE_PDF_REPAYMENT_PLAN = "Otplatni_plan.pdf"

# --- Pre-extracted Data for MVP ---
# (Based on analysis of provided PDFs: DocumentServlet.pdf, FileDownloadServlet.pdf, Otplatni_plan.pdf)
extracted_data = {
    "dodatak_datum_zakljucenja": "24.02.2025.",
    "dodatak_mjesto_zakljucenja": "ĐAKOVO",
    "dodatak_broj": "1",
    "korisnik_ime_prezime": "BARBARA STAŽIĆ",
    "korisnik_adresa": "JOSIPA JURJA STROSSMAYERA 129, 31000 OSIJEK",
    "korisnik_oib": "61897713961",
    "original_ugovor_naziv": "UGOVOR O NENAMJENSKOM GOTOVINSKOM KREDITU",
    "original_ugovor_broj_partije": "9919479387",
    "original_ugovor_datum_zakljucenja": "18.03.2024.",
    "original_ugovor_iznos": 11000.00,
    "original_ugovor_valuta": "EUR",
    "uplata_iznos": 3000.00, # Calculated: 9158.10 (from DocumentServlet) - 6158.10 (from Otplatni_plan)
    "nova_glavnica_iznos": 6158.10,
    "novi_anuitet_iznos": 185.26,
    "solidarni_duznik_ime_prezime": None, # Not found in docs
    "solidarni_duznik_adresa": None,
    "solidarni_duznik_oib": None,
    "solidarni_jamac_ime_prezime": None, # Not found in docs
    "solidarni_jamac_adresa": None,
    "solidarni_jamac_oib": None,
    # Add more fields as needed for other placeholders if missed
}

# --- Helper Function for Croatian Number Words ---
def num_to_hr_words(number, currency="EUR"):
    """Converts a number to Croatian words, handling euros and cents."""
    try:
        # num2words should be imported at the top level
        pass
    except NameError: # If num2words failed import at top
        return f"[Greška: num2words nedostaje {number} {currency}]"


    if not isinstance(number, (int, float)):
        return str(number) # Return as string if not a number

    integer_part = int(number)
    fractional_part = round((number - integer_part) * 100)

    # Use feminine forms for amounts if appropriate (more complex grammar)
    # For MVP, default masculine/neuter forms from num2words are used.
    words = num2words(integer_part, lang='hr') # num2words directly accessible

    if currency == "EUR":
        # Handle declension for 'euro' based on number
        if integer_part % 10 == 1 and integer_part % 100 != 11:
            words += " euro"
        elif 2 <= integer_part % 10 <= 4 and not (12 <= integer_part % 100 <= 14):
             words += " eura"
        else:
             words += " eura" # Default/genitive plural

        if fractional_part > 0:
            words += " i " + num2words(fractional_part, lang='hr')
            # Handle declension for 'cent'
            if fractional_part % 10 == 1 and fractional_part % 100 != 11:
                 words += " cent"
            elif 2 <= fractional_part % 10 <= 4 and not (12 <= fractional_part % 100 <= 14):
                 words += " centa"
            else:
                 words += " centi" # Default/genitive plural

    elif currency == "HRK": # Example if needed later
         # Handle 'kuna' declension similarly if needed
         words += " kuna"
         if fractional_part > 0:
            words += " i " + num2words(fractional_part, lang='hr') + " lipa" # Add lipa declension if needed
    else: # Handle non-currency numbers or other currencies
        if fractional_part > 0:
             # Generic decimal representation
             words += " cijelih " + num2words(fractional_part, lang='hr') + " stotinki"

    return words

def date_to_hr_words(date_str):
    """ Converts a DD.MM.GGGG date string to Croatian words (genitive form). """
    try:
        # num2words should be imported at the top level
        dt_obj = datetime.strptime(date_str, "%d.%m.%Y")
        day = dt_obj.day
        month = dt_obj.month
        year = dt_obj.year

        # Ordinal day requires genitive case for "day month year" structure
        # num2words 'to=ordinal' gives nominative (prvi, drugi...)
        # We need genitive (prvog, drugog...)
        day_ord_nom = num2words(day, lang='hr', to='ordinal')
        if day_ord_nom.endswith('i'): # prvi, drugi, treći...
            day_ord_gen = day_ord_nom[:-1] + 'og' # prvog, drugog, trećeg... (approximation)
        elif day_ord_nom.endswith('a'): # četvrta...
             day_ord_gen = day_ord_nom[:-1] + 'e' # četvrte
        else:
            day_ord_gen = day_ord_nom # Fallback for other cases

        months_genitive = [
            "siječnja", "veljače", "ožujka", "travnja", "svibnja", "lipnja",
            "srpnja", "kolovoza", "rujna", "listopada", "studenog", "prosinca"
        ]
        month_words = months_genitive[month - 1]
        year_words = num2words(year, lang='hr') # Year in cardinal words

        return f"{day_ord_gen} {month_words} {year_words}"

    except Exception as e:
        print(f"WARN: Error converting date {date_str} to words: {e}")
        return f"[Datum Greška: {date_str}]"


# --- Pydantic Model for the Tool Input ---
class FillDocxArgs(BaseModel):
    template_file_path: str = Field(..., description="Relative path to the DOCX template file within the MCP filesystem.")
    output_file_path: str = Field(..., description="Relative path for the generated DOCX file within the MCP filesystem.")
    data: dict = Field(..., description="Dictionary containing data to fill into the template placeholders.")

# --- Python Code to run via MCP ---
# This code will be executed by the '@pydantic/mcp-run-python' server
# It needs access to 'python-docx' and 'num2words' libraries.

PYTHON_CODE_FOR_MCP = """
import sys
import json
import re
from datetime import datetime

# Attempt to import required libraries within MCP context
try:
    from docx import Document
    from docx.shared import Pt
    # from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Uncomment if needed
except ImportError:
    print(json.dumps({"success": False, "error": "python-docx not found in MCP environment"}), file=sys.stderr)
    sys.exit(1)
try:
    import num2words as nw # Use alias to avoid potential conflicts
except ImportError:
    print(json.dumps({"success": False, "error": "num2words not found in MCP environment"}), file=sys.stderr)
    # Define fallback functions if needed, or just fail. Failing is better for required libs.
    sys.exit(1)

# --- Helper Functions (copied within MCP execution context) ---
def num_to_hr_words(number, currency="EUR"):
    if not isinstance(number, (int, float)): return str(number)
    integer_part = int(number)
    fractional_part = round((number - integer_part) * 100)
    words = nw.num2words(integer_part, lang='hr')
    if currency == "EUR":
        if integer_part % 10 == 1 and integer_part % 100 != 11: words += " euro"
        elif 2 <= integer_part % 10 <= 4 and not (12 <= integer_part % 100 <= 14): words += " eura"
        else: words += " eura"
        if fractional_part > 0:
            words += " i " + nw.num2words(fractional_part, lang='hr')
            if fractional_part % 10 == 1 and fractional_part % 100 != 11: words += " cent"
            elif 2 <= fractional_part % 10 <= 4 and not (12 <= fractional_part % 100 <= 14): words += " centa"
            else: words += " centi"
    # Add other currency logic if needed (e.g., HRK)
    else: words += f" {currency}"
    return words

def date_to_hr_words(date_str):
    try:
        dt_obj = datetime.strptime(date_str, "%d.%m.%Y")
        day, month, year = dt_obj.day, dt_obj.month, dt_obj.year
        months_genitive = ["siječnja", "veljače", "ožujka", "travnja", "svibnja", "lipnja",
                           "srpnja", "kolovoza", "rujna", "listopada", "studenog", "prosinca"]
        day_ord_nom = nw.num2words(day, lang='hr', to='ordinal')
        day_ord_gen = day_ord_nom[:-1] + 'og' if day_ord_nom.endswith('i') else day_ord_nom # Basic genitive adjustment
        return f"{day_ord_gen} {months_genitive[month - 1]} {nw.num2words(year, lang='hr')}"
    except Exception as e:
        print(f"WARN (MCP): Error converting date {date_str} to words: {e}", file=sys.stderr)
        return f"[Datum Greška: {date_str}]"

# --- Main Fill Logic (within MCP execution context) ---
def fill_template(template_path, output_path, data):
    # print(f"MCP: Filling template: {template_path}") # Debug print
    # print(f"MCP: Output path: {output_path}") # Debug print
    # print(f"MCP: Data: {data}") # Debug print (be careful with sensitive data)

    try:
        doc = Document(template_path)
    except Exception as e:
        return {"success": False, "error": f"MCP Failed to open template '{template_path}': {e}"}

    # --- Placeholder Replacement Logic ---
    # Use paragraph analysis for more robust replacement than simple run replace

    in_article_1 = False
    in_article_2 = False
    article_1_date_replaced = False
    article_1_amount_num_replaced = False
    article_1_amount_words_replaced = False
    article_2_payment_num_replaced = False
    article_2_payment_words_replaced = False
    article_2_principal_num_replaced = False
    article_2_principal_words_replaced = False
    article_2_annuity_num_replaced = False
    article_2_annuity_words_replaced = False
    article_2_clause_to_remove_found = False
    article_2_clause_to_keep_found = False

    paragraphs_to_remove = [] # Store paragraphs to be deleted later

    # --- General Replacements (Header/Footer/Generic) ---
    # These can be done first if they are less context-dependent
    generic_replacements = {
        # Use raw strings for regex patterns
        r'\\[DD\\.MM\\.GGGG\\.\\]': data.get('dodatak_datum_zakljucenja', '[DATUM_DODATKA]'),
        r'\\[upisati mjesto\\]': data.get('dodatak_mjesto_zakljucenja', '[MJESTO_DODATKA]'),
        r'Dodatak br\.? __?': f"Dodatak br. {data.get('dodatak_broj', 'X')}", # Variations
        r'Dodatka br\.? __?': f"Dodatka br. {data.get('dodatak_broj', 'X')}",
        r'Dodatkom br\.? __?': f"Dodatkom br. {data.get('dodatak_broj', 'X')}",
        r'Dodatku br\.? __?': f"Dodatku br. {data.get('dodatak_broj', 'X')}",
        r'\\[upisati naziv ugovora.*?kreditu\\]': data.get('original_ugovor_naziv', '[NAZIV_UGOVORA]'),
        r'\\[?9910000000\\]?': data.get('original_ugovor_broj_partije', '[BROJ_PARTIJE]'),
        r'IME I PREZIME': data.get('korisnik_ime_prezime', '[IME_PREZIME_KORISNIKA]'), # Assumes first occurrence
        r'Adresa': data.get('korisnik_adresa', '[ADRESA_KORISNIKA]'), # Assumes first occurrence
        r'_{8,}': data.get('korisnik_oib', '[OIB_KORISNIKA]'), # Assumes first OIB placeholder
        r'\\[upisati datum slovima\\]': date_to_hr_words(data.get('dodatak_datum_zakljucenja', '')), # Addendum date words
    }

    for para in doc.paragraphs:
         original_text = para.text
         modified_text = original_text
         for pattern, replacement in generic_replacements.items():
              modified_text = re.sub(pattern, str(replacement), modified_text, count=1) # Replace once per pattern per para initially

         if modified_text != original_text:
              # Clear existing runs and add the modified text
              # Warning: This destroys formatting within the paragraph!
              # A run-level replacement is better for preserving formatting if possible.
              # For MVP, let's try simple paragraph text update first.
              para.text = modified_text


    # --- Context-Specific Replacements & Logic ---
    for i, para in enumerate(doc.paragraphs):
        text = para.text # Get potentially modified text from previous step

        # Track current article
        if text.strip().startswith("Članak 1."):
            in_article_1 = True; in_article_2 = False
        elif text.strip().startswith("Članak 2."):
            in_article_1 = False; in_article_2 = True
        elif text.strip().startswith("Članak 3."):
            in_article_1 = False; in_article_2 = False

        # --- Article 1 ---
        if in_article_1:
            if "dana " in text and not article_1_date_replaced: # More specific context needed if "[DD.MM.GGGG.]" was already replaced generically
                # Assume generic replacements handled date/place/addendum# already.
                # We need to replace the specific date/words for the *original* contract here.
                # Requires more unique placeholders in template, e.g., [DATUM_ORIG_UGOVORA], [DATUM_ORIG_UGOVORA_SLOVIMA]
                # Workaround: Rely on order or search for specific surrounding text if template is fixed.
                # For now, let's assume generic replacement worked for the *first* date found.
                # Need to handle ORIGINAL contract date/amount specifically if different placeholders aren't used.
                 pass # Skip specific replacement if generic handled it, needs template refinement

            # Remove conditional clauses if needed
            if "Nova hrvatska banka dioničko društvo" in text: pass # Keep intro for now
            if "nakon uvođenja EUR-a kao službene valute" in text: paragraphs_to_remove.append(para)
            if "nakon pripajanja pravnog prednika Banke Banci" in text: paragraphs_to_remove.append(para)

        # --- Article 2 ---
        if in_article_2:
             # Replace amounts and words, ensure correct placeholders are targeted
             # Using more specific markers if possible:
             if "u iznosu od [XX.XXX,XX] EUR" in text and not article_2_payment_num_replaced:
                  amount_str = f"{data.get('uplata_iznos', 0.0):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                  para.text = para.text.replace("[XX.XXX,XX] EUR", f"{amount_str} EUR")
                  article_2_payment_num_replaced = True
                  if "(slovima: [upisati slovima iznos])" in text and not article_2_payment_words_replaced:
                      para.text = para.text.replace("[upisati slovima iznos]", num_to_hr_words(data.get('uplata_iznos', 0.0), "EUR"))
                      article_2_payment_words_replaced = True

             if "ista sada iznosi [XX.XXX,XX] EUR" in text and not article_2_principal_num_replaced:
                  amount_str = f"{data.get('nova_glavnica_iznos', 0.0):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                  para.text = para.text.replace("[XX.XXX,XX] EUR", f"{amount_str} EUR")
                  article_2_principal_num_replaced = True
                  if "(slovima: [upisati slovima iznos])" in text and not article_2_principal_words_replaced:
                       para.text = para.text.replace("[upisati slovima iznos]", num_to_hr_words(data.get('nova_glavnica_iznos', 0.0), "EUR"))
                       article_2_principal_words_replaced = True

             # Annuity clauses
             shorter_term_marker = "mijenja datum dospijeća zadnjeg anuiteta"
             # Assume the clause *without* the marker is the one to remove if both are present
             is_shorter_term_clause = shorter_term_marker in text
             is_potentially_same_term_clause = text.strip().startswith("2)") and "[XX.XXX,XX] EUR" in text and not is_shorter_term_clause

             if is_shorter_term_clause and not article_2_annuity_num_replaced:
                  amount_str = f"{data.get('novi_anuitet_iznos', 0.0):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                  para.text = para.text.replace("[XX.XXX,XX] EUR", f"{amount_str} EUR")
                  article_2_annuity_num_replaced = True
                  if "(slovima: [upisati slovima iznos])" in text and not article_2_annuity_words_replaced:
                       para.text = para.text.replace("[upisati slovima iznos]", num_to_hr_words(data.get('novi_anuitet_iznos', 0.0), "EUR"))
                       article_2_annuity_words_replaced = True
                  article_2_clause_to_keep_found = True

             elif is_potentially_same_term_clause and not article_2_clause_to_remove_found:
                  # Mark the paragraph containing the alternative clause for removal
                  paragraphs_to_remove.append(para)
                  article_2_clause_to_remove_found = True


    # --- Remove marked paragraphs ---
    for para in reversed(paragraphs_to_remove):
        try:
            p_element = para._element
            p_element.getparent().remove(p_element)
        except Exception as e:
            print(f"WARN (MCP): Could not remove paragraph: {e}", file=sys.stderr)

    # --- Final Save ---
    try:
        doc.save(output_path)
        # print(f"MCP: Successfully generated document: {output_path}") # Debug
        return {"success": True, "output_path": output_path}
    except Exception as e:
        return {"success": False, "error": f"MCP Failed to save document '{output_path}': {e}"}

# --- Entry point for MCP execution ---
if __name__ == '__main__':
    # Expecting args as a JSON string from stdin
    args_json = sys.stdin.read()
    try:
        args = json.loads(args_json)
        result = fill_template(
            template_path=args['template_file_path'],
            output_path=args['output_file_path'],
            data=args['data']
        )
        print(json.dumps(result)) # Output result as JSON to stdout
    except Exception as e:
        # Output error as JSON to stderr
        print(json.dumps({"success": False, "error": f"MCP Script Error: {e}"}), file=sys.stderr)
        sys.exit(1)
"""

# --- Agent Setup ---

# Define MCP Servers
servers = [
    MCPServerStdio(
        'npx',
        # Ensure npx and the MCP packages are available in the Docker container's PATH
        ['-y', '@pydantic/mcp-run-python', 'stdio', '--python-path', sys.executable], # Use current python env
        
    ),
    MCPServerStdio(
        'npx',
        ['-y', '@modelcontextprotocol/server-filesystem', str(BASE_WORKING_DIR)], # Serve files from BASE_WORKING_DIR
       
    )
]

# Define the AI Model
model = OpenAIModel(
    'gpt-4o', # Or another suitable model like gpt-3.5-turbo
    provider=OpenAIProvider(api_key=OPENAI_API_KEY)
)

# Create the Agent
agent = Agent(
    model=model,
    system_prompt="You are a helpful banking assistant. Your task is to fill document templates using provided data and tools. You can execute Python code and access a local filesystem.",
    mcp_servers=servers,
    # debug=True # Enable for more verbose output from Pydantic AI
)

# --- Main Execution Logic ---
async def main():
    print("Pozdrav! Što mogu učiniti za vas danas?") # Greeting message
    print("Unesite 'bok' ili 'doviđenja' za izlaz.") # Instructions to exit
    print("=======================================")

    try:
        # Start MCP servers once and keep them running for the loop
        async with agent.run_mcp_servers():
            while True: # Start infinite loop for conversation
                try:
                    user_input = input("\n[Vi] ").strip() # Get user input
                except EOFError:
                    # Handle case where input stream is closed unexpectedly
                    print("\n[Asistent] Nema više unosa. Doviđenja!")
                    break

                # Check for exit commands
                exit_commands = ["bok", "doviđenja", "dovidenja", "exit", "quit", "bye"]
                if user_input.lower() in exit_commands:
                    print("[Asistent] Doviđenja!")
                    break # Exit the loop

                # --- If not exiting, assume user wants to generate the document ---
                print("[Asistent] U redu, pokrećem generiranje dokumenta...")

                # Define paths relative to BASE_WORKING_DIR for MCP
                template_path_mcp = TEMPLATE_FILENAME
                output_path_mcp = OUTPUT_FILENAME

                # Prepare the arguments for the Python code execution tool
                fill_args = FillDocxArgs(
                    template_file_path=template_path_mcp,
                    output_file_path=output_path_mcp,
                    data=extracted_data # Using pre-extracted data for MVP
                )
                # Use model_dump_json for Pydantic v2+
                fill_args_json = fill_args.model_dump_json()


                # Construct the prompt for the agent to use the Python executor
                # Simplified prompt focusing on execution
                prompt = f"""
                Execute the python code using the 'python_executor' tool.
                Pass the following JSON string as the argument to the code:
                {fill_args_json}

                Report the JSON result provided by the tool.
                The python code to be executed is:
                ```python
                {PYTHON_CODE_FOR_MCP}
                ```
                """

                print("\n[Agent Task] Asking agent to execute Python code via MCP...")

                # Run the agent task for document generation
                # Pass the Python code itself as part of the context/prompt
                result = await agent.run(prompt)

                print("\n[Agent Result Raw Data]")
                print(result.data) # Agent's textual response, hopefully includes JSON from tool

                # Attempt to parse JSON output from the agent's response
                tool_result = None
                try:
                    # Look for JSON in the agent's output (this might need refinement)
                    json_match = re.search(r'\{.*\}', result.data, re.DOTALL)
                    if json_match:
                        tool_result = json.loads(json_match.group())
                        print("\n[Parsed Tool Result]")
                        print(tool_result)
                    else:
                        print("\n[Verification] Could not find JSON result in agent output.")
                except json.JSONDecodeError:
                     print("\n[Verification] Could not parse JSON from agent output.")
                except Exception as parse_err:
                     print(f"\n[Verification] Error parsing agent output: {parse_err}")


                # Manual check based on expected output path
                output_file_full_path = BASE_WORKING_DIR / OUTPUT_FILENAME
                if output_file_full_path.exists():
                     print(f"\n[Verification] SUCCESS: Output file '{output_file_full_path}' found.")
                     # Optional: Check modification time to be recent
                else:
                     print(f"\n[Verification] WARNING: Output file '{output_file_full_path}' was NOT found.")
                     if tool_result and not tool_result.get("success"):
                         print(f"  Tool reported error: {tool_result.get('error')}")

                print("\n[Asistent] Zadatak generiranja dokumenta je završen (provjerite verifikaciju iznad). Što sljedeće želite učiniti?")
                # Loop continues, waits for next input

    except Exception as e:
        print(f"\nFATAL ERROR during agent execution: {e}")
        import traceback
        traceback.print_exc() # Print full traceback for debugging
    finally:
        # This block executes when the loop breaks or an error occurs outside the `async with`
        print("\nZavršavam s radom.")

# --- Final execution block ---
if __name__ == '__main__':
    # Ensure the working directory exists
    if not BASE_WORKING_DIR.exists():
         print(f"ERROR: Base working directory '{BASE_WORKING_DIR}' does not exist. Please create it or update the path in the script or via AGENT_WORKING_DIR env var.")
         sys.exit(1)
    # Ensure the template file exists
    template_full_path = BASE_WORKING_DIR / TEMPLATE_FILENAME
    if not template_full_path.exists():
         print(f"ERROR: Template file '{TEMPLATE_FILENAME}' not found in '{BASE_WORKING_DIR}'.")
         print(f"Full path checked: {template_full_path}")
         sys.exit(1)

    # Check Python executable permissions (relevant in some Docker setups)
    if not os.access(sys.executable, os.X_OK):
         print(f"WARN: Python executable '{sys.executable}' might not have execute permissions.")

    print(f"--- Starting Agent ---")
    print(f"Python Executable: {sys.executable}")
    print(f"Working Directory: {BASE_WORKING_DIR}")
    print(f"Template Path: {template_full_path}")
    print(f"----------------------")

    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\nPrimljen prekid (Ctrl+C). Završavam.")